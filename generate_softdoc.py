from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def add_code_block(doc, title, filepath):
    doc.add_heading(title, level=2)
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line.rstrip('\n'))
        run.font.name = 'Consolas'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Pt(12)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05

# 新建文档
doc = Document()
doc.add_heading('轨道交通车站信息汇报助手 源代码（软著申报用）', 0)

# 目录
doc.add_paragraph('目录：\n1. index.html\n2. script.js\n3. styles.css')

# 添加各文件内容
add_code_block(doc, '【index.html】', 'index.html')
add_code_block(doc, '【script.js】', 'script.js')
add_code_block(doc, '【styles.css】', 'styles.css')

# 保存
doc.save('轨道交通车站信息汇报助手-源代码软著版.docx')
import os
print('当前目录：', os.getcwd())
print('Word文档已生成：轨道交通车站信息汇报助手-源代码软著版.docx')